import os
from typing import List, Dict, Any
from loguru import logger

try:
    from pypdf import PdfReader
    from docx import Document
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document as LangDocument
    import chardet
except ImportError as e:
    logger.error(f"Import error: {e}")
    raise


class DocumentProcessor:
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):  # 减小块大小和重叠
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )

    def extract_text(self, file_path: str, file_type: str) -> str:
        """提取文档文本内容"""
        try:
            file_size = os.path.getsize(file_path) / 1024 / 1024  # MB
            logger.info(f"处理文件: {file_path}, 大小: {file_size:.2f}MB, 类型: {file_type}")

            if file_size > 50:  # 大于50MB的文件
                logger.warning(f"大文件警告: {file_path} ({file_size:.2f}MB)")

            if file_type == 'pdf':
                return self._extract_pdf(file_path, file_size)
            elif file_type == 'docx':
                return self._extract_docx(file_path, file_size)
            elif file_type in ['md', 'txt']:
                return self._extract_text_file(file_path, file_size)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise

    def _extract_pdf(self, file_path: str, file_size: float) -> str:
        """优化的大PDF文件处理"""
        reader = PdfReader(file_path)
        total_pages = len(reader.pages)
        text = ""

        logger.info(f"PDF文件共 {total_pages} 页")

        # 对于大文件，分批处理避免内存溢出
        batch_size = 50  # 每批处理50页
        processed_pages = 0

        for start_page in range(0, total_pages, batch_size):
            end_page = min(start_page + batch_size, total_pages)
            batch_text = ""

            for page_num in range(start_page, end_page):
                try:
                    page = reader.pages[page_num]
                    page_text = page.extract_text()

                    # 清理PDF文本中的常见问题
                    cleaned_text = self._clean_pdf_text(page_text)
                    if cleaned_text.strip():  # 只添加非空文本
                        batch_text += cleaned_text + "\n"

                except Exception as e:
                    logger.warning(f"第 {page_num + 1} 页提取失败: {e}")
                    continue

            text += batch_text
            processed_pages = end_page
            logger.info(f"已处理 {processed_pages}/{total_pages} 页")

            # 对于超大文件，每处理完一个批次就进行分块，避免内存积累
            if file_size > 20 and len(text) > 100000:  # 大于20MB且文本超过10万字符
                logger.info("大文件中间分块处理...")
                # 这里可以添加中间处理逻辑

        logger.info(f"PDF文本提取完成，总字符数: {len(text)}")
        return text

    def _clean_pdf_text(self, text: str) -> str:
        """清理PDF提取的文本"""
        if not text:
            return ""

        # 移除过多的空白字符
        import re
        text = re.sub(r'\s+', ' ', text)

        # 移除页眉页脚等常见噪音
        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            line = line.strip()
            # 跳过可能是页码或页眉的内容
            if (len(line) < 5 and line.isdigit()) or \
                    any(keyword in line.lower() for keyword in ['chapter', 'section', 'page']):
                continue
            if line:
                cleaned_lines.append(line)

        return '\n'.join(cleaned_lines)

    def _extract_docx(self, file_path: str, file_size: float) -> str:
        """优化的大DOCX文件处理"""
        doc = Document(file_path)
        text = ""

        # 分批处理段落
        paragraphs_processed = 0
        total_paragraphs = len(doc.paragraphs)

        logger.info(f"DOCX文件共 {total_paragraphs} 个段落")

        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                text += paragraph.text + "\n"
                paragraphs_processed += 1

            # 每处理100个段落输出进度
            if i % 100 == 0 and i > 0:
                logger.info(f"已处理 {i}/{total_paragraphs} 个段落")

        logger.info(f"DOCX文本提取完成，总段落数: {paragraphs_processed}")
        return text

    def _extract_text_file(self, file_path: str, file_size: float) -> str:
        """优化的大文本文件处理"""
        # 检测文件编码
        encoding = self._detect_encoding(file_path)

        # 对于大文本文件，使用流式读取
        if file_size > 10:  # 大于10MB的文本文件
            return self._stream_read_large_text(file_path, encoding)
        else:
            # 小文件直接读取
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                content = f.read()
            return content

    def _detect_encoding(self, file_path: str) -> str:
        """检测文件编码"""
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)  # 读取前10KB检测编码

            detected = chardet.detect(raw_data)
            encoding = detected.get('encoding', 'utf-8')
            confidence = detected.get('confidence', 0)

            logger.info(f"检测到编码: {encoding} (置信度: {confidence:.2f})")

            # 优先使用中文友好的编码
            if encoding.lower() in ['gb2312', 'gbk', 'gb18030']:
                return 'gb18030'  # 最全的中文编码
            elif encoding.lower() == 'ascii':
                return 'utf-8'  # ASCII兼容UTF-8
            else:
                return encoding

        except Exception as e:
            logger.warning(f"编码检测失败: {e}, 使用UTF-8")
            return 'utf-8'

    def _stream_read_large_text(self, file_path: str, encoding: str) -> str:
        """流式读取大文本文件"""
        content = ""
        chunk_size = 8192  # 8KB chunks
        total_size = os.path.getsize(file_path)
        bytes_read = 0

        logger.info(f"流式读取大文本文件: {file_path}")

        try:
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                while True:
                    chunk = f.read(chunk_size)
                    if not chunk:
                        break
                    content += chunk
                    bytes_read += len(chunk.encode(encoding))

                    # 每读取1MB输出进度
                    if bytes_read % (1024 * 1024) < chunk_size:
                        progress = (bytes_read / total_size) * 100
                        logger.info(f"读取进度: {progress:.1f}% ({bytes_read // 1024}KB/{total_size // 1024}KB)")

        except Exception as e:
            logger.error(f"流式读取失败: {e}")
            raise

        return content

    def split_documents(self, text: str, metadata: Dict[str, Any] = None) -> List[LangDocument]:
        """将文本分割成块，优化大文本处理"""
        if metadata is None:
            metadata = {}

        logger.info(f"开始分割文本，总长度: {len(text)} 字符")

        # 对于超长文本，可以考虑分批分割
        if len(text) > 1000000:  # 超过100万字符
            logger.warning("文本过长，采用分批分割策略")
            return self._split_large_text_in_batches(text, metadata)
        else:
            documents = self.text_splitter.create_documents([text], [metadata])
            logger.info(f"分割完成: {len(documents)} 个块")
            return documents

    def _split_large_text_in_batches(self, text: str, metadata: Dict[str, Any]) -> List[LangDocument]:
        """分批处理超长文本"""
        batch_size = 200000  # 每批20万字符
        all_documents = []
        total_batches = (len(text) + batch_size - 1) // batch_size

        for i in range(total_batches):
            start_idx = i * batch_size
            end_idx = min((i + 1) * batch_size, len(text))
            batch_text = text[start_idx:end_idx]

            logger.info(f"处理批次 {i + 1}/{total_batches}")
            batch_docs = self.text_splitter.create_documents([batch_text], [metadata])
            all_documents.extend(batch_docs)

        logger.info(f"分批分割完成，总共 {len(all_documents)} 个块")
        return all_documents